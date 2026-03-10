import os
from PyQt5.QtCore import QSize
from qgis.core import QgsProject, QgsFeature
from PyQt5.QtGui import QImage, QPainter
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Define output path (Desktop for easy access)
output_path = os.path.expanduser("~/Desktop/output_coft_mac.docx")

# Define Layer Names
layer1_name = 'Barrierefreihe_Elemente_POIs_Sandweiler'
layer2_name = 'Erfassungsboegen_V3_QGIS_Join_PMR-Sandweiler_08012025'
pois_layer_name = 'POIs_total_Sandweiler'

# Get layers from QGIS project
layer1 = QgsProject.instance().mapLayersByName(layer1_name)[0]
layer2 = QgsProject.instance().mapLayersByName(layer2_name)[0]
pois_layer = QgsProject.instance().mapLayersByName(pois_layer_name)[0]

# Initialize Word Document
doc = Document()

# Set document to A4 landscape format
section = doc.sections[0]
section.page_width = Cm(29.7)  # A4 width
section.page_height = Cm(21.0)  # A4 height
section.orientation = 1  # Landscape orientation

# Define category mapping function
def get_kategorie_name(kategorie):
    kategorie_mapping = {
        1: "Orientierung", 2: "Haupteingang", 3: "Nebeneingang", 4: "Rampe",
        5: "Treppe", 6: "Parkplatz", 7: "Zuwegung", 8: "Einfahrt",
        9: "Fußgängerübergang", 10: "Bushaltestelle", 11: "Bürgersteig", 12: "PMR-Parkplatz"
    }
    return kategorie_mapping.get(kategorie, "Unknown")

# Zustand Mapping
zustand_mapping = {'1': 'gut', '2': 'mittel', '3': 'schlecht', '4': 'nicht vorhanden (keine bewertung)'}

# Iterate through POIs
for poi_feature in pois_layer.getFeatures():
    poi_name = poi_feature['POI_Name']
    poi_address = poi_feature['POI_Adresse']

    # Add heading for POI
    heading = doc.add_heading(poi_name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(34)

    doc.add_paragraph(poi_address).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Function to extract category
    def extract_kategorie(element_id):
parts = element_id.split('_')
        return int(parts[-2]) if len(parts) > 1 else -1

    # Group elements by category
    element_groups = {}
    for feature in layer1.getFeatures():
        element_id = feature['Element_ID']
        kategorie = extract_kategorie(element_id)
        element_groups.setdefault(kategorie, []).append(feature)

    for kategorie in sorted(element_groups.keys()):
        features = element_groups[kategorie]
        element_ids = [f['Element_ID'] for f in features]
        kategorie_name = get_kategorie_name(kategorie)

        doc.add_heading(f"Kategorie: {kategorie} - {kategorie_name}", level=2)
        doc.add_paragraph(f"Represented by Element_ID(s): {', '.join(map(str, element_ids))}")

        for feature in features:
            element_id = feature['Element_ID']
            lage_autom = feature['Lage_autom']
            nummer = feature['Nummer']
            fid = feature['fid']
            nearest_poi_id = feature['Nearest_POI_ID']
            nearest_poi_name = feature['Nearest_POI_Name']
            nearest_poi_distance = feature['Nearest_POI_distance_(m2)']
            foto_element_export = feature['foto_element_export']

            doc.add_heading(f"Element_ID {element_id}", level=3)
            doc.add_paragraph(f"Lage: {lage_autom}")
            doc.add_paragraph(f"Kategorie Name: {kategorie_name}")
            doc.add_paragraph(f"Nummer: {nummer}")
            doc.add_paragraph(f"FID: {fid}")
            doc.add_paragraph(f"Nearest_POI_ID: {nearest_poi_id}")
            doc.add_paragraph(f"Nearest_POI_Name: {nearest_poi_name}")
            doc.add_paragraph(f"Distance from POI: {nearest_poi_distance} meters" if nearest_poi_distance else "Located within POI.")

            # Check for images and insert them
            if foto_element_export and os.path.exists(foto_element_export):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(foto_element_export, width=Inches(4.5))

            # Get survey data
            matching_features = [feature2 for feature2 in layer2.getFeatures() if feature2['Element_ID'] == element_id]

            if matching_features:
                doc.add_heading("Erfassungsboegen", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Create table
                table = doc.add_table(rows=1 + len(matching_features), cols=6)
                table.style = 'Table Grid'
                
                # Force table to not autofit, so custom widths apply
                table.autofit = False


                # Define custom column widths
                column_widths = [Cm(1.5), Cm(4), Cm(7), Cm(4), Cm(4), Cm(5)]  

                # Define headers
                headers = ['Nr.', 'Merkmal', 'Beschreibung', 'Zustand', 'Foto', 'Hinweis/Maße']

                # Apply headers and column widths
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.width = column_widths[i]
                
                # Define cell colors (Hex codes)
                zustand_colors = {
                    '1': 'CCFFCC',  # Green
                    '2': 'FFFF99',  # Yellow
                    '3': 'FF6666',  # Red
                    '4': 'ADD8E6'   # Light Blue
                }

                # Populate table rows with data
                for row_idx, feature2 in enumerate(matching_features):
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = str(feature2['Nr.'])
                    row_cells[1].text = str(feature2['Merkmal'])
                    row_cells[2].text = str(feature2['Beschreibung'])
                    row_cells[3].text = zustand_mapping.get(str(feature2['Zustand: 1 (gut) - 3 (schlecht)']), ' ')
                    row_cells[4].text = str(feature2['Fotoname'])
                    row_cells[5].text = str(feature2['Hinweis/Maße'])
# Apply column widths for data rows
                    for i, cell in enumerate(row_cells):
                        cell.width = column_widths[i]
                
                    # Get Zustand value from the feature2 dictionary
                    zustand_value = str(feature2['Zustand: 1 (gut) - 3 (schlecht)']) 
                    
                    #color for specific column
                    if zustand_value in zustand_colors:
                        color = zustand_colors[zustand_value]
                        cell_element = row_cells[3]._element.get_or_add_tcPr()
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                        cell_element.append(shading)

#color for entire row
"""                for row_idx, feature2 in enumerate(matching_features):
                    row_cells = table.rows[row_idx + 1].cells
                    zustand_value = str(feature2['Zustand: 1 (gut) - 3 (schlecht)'])
                    color_mapping = {'1': 'CCFFCC', '2': 'FFFFCC', '3': 'FF0000', '4': 'ADD8E6'}
                    row_color = color_mapping.get(zustand_value, 'FFFFFF')  
                    
                    for cell in row_cells:
                        cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{row_color}"/>'))
"""

doc.add_paragraph("\n")

doc.add_page_break()
#you can adjust the size of the table here
 # Adjust column alignment and width
widths = [0.7, 2.0, 2.5, 1.2, 1.5, 2.5]  
for i, cell in enumerate(row_cells):
  	cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
cell.width = Inches(widths[i])

# Add page numbers
section = doc.sections[-1]
footer = section.footer
paragraph = footer.paragraphs[0]
field_code = '<w:fldSimple {0} w:instr="PAGE" w:result="1"/>'.format(nsdecls('w'))
paragraph._element.append(parse_xml(field_code))
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save(output_path)
print(f"Word document successfully saved at: {output_path}")

